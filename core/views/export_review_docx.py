from __future__ import annotations

import os
from io import BytesIO

from django.contrib.auth.decorators import login_required
from django.http import Http404, HttpResponse, HttpResponseForbidden
from django.utils import timezone

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips

from core.models import MedicatieReviewAfdeling, MedicatieReviewPatient
from core.views._helpers import _static_abs_path, can
from core.views.export_review_pdf import PdfPatientBlock, _build_patient_block


# ── Colour palette ─────────────────────────────────────────────────────────────
COL_HEADER_BG = "EEF2F7"
COL_SECTION_FG = "0C2742"
COL_MUTED       = "334155"  # dark slate for secondary text
COL_DIVIDER = "D1D5DB"
COL_COMMENT_BG = "F3F6FA"

PHARMACY_NAME = "Apotheek Jansen Amersfoort"
PHARMACY_EMAIL = "instellingen@apotheekjansen.com"

PREPARED_BY_EMAIL_FIXED = PHARMACY_EMAIL

DOC_WIDTH_DXA = 9072
HEADER_LOGO_COL_DXA = 900   # just wider than the 1.5 cm logo image
HEADER_INFO_COL_DXA = DOC_WIDTH_DXA - HEADER_LOGO_COL_DXA


# ── XML helpers ────────────────────────────────────────────────────────────────

def _get_or_create_tblPr(tbl):
    tbl_el = tbl._tbl
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)
    return tblPr


def _set_tbl_width(tbl, width_dxa: int = DOC_WIDTH_DXA) -> None:
    tblPr = _get_or_create_tblPr(tbl)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(width_dxa))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _set_tbl_fixed_layout(tbl) -> None:
    tblPr = _get_or_create_tblPr(tbl)
    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def _set_tbl_grid(tbl, col_widths_dxa: list) -> None:
    """
    Write explicit tblGrid so fixed-layout column widths are respected.
    python-docx generates a default single-column grid that ignores _set_col_width.
    """
    tbl_el = tbl._tbl
    for old in tbl_el.findall(qn("w:tblGrid")):
        tbl_el.remove(old)
    grid = OxmlElement("w:tblGrid")
    for w in col_widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    # tblGrid must follow tblPr in schema order
    tblPr = tbl_el.find(qn("w:tblPr"))
    tblPr.addnext(grid)


def _set_table_grid_style(tbl) -> None:
    """
    Set tblStyle to TableGrid via direct XML — bypasses python-docx style lookup.
    TableGrid is a Word built-in that guarantees borders exist as a base; our
    per-cell _cell_borders calls then override the colour.
    """
    tblPr = _get_or_create_tblPr(tbl)
    for old in tblPr.findall(qn("w:tblStyle")):
        tblPr.remove(old)
    tblStyle = OxmlElement("w:tblStyle")
    tblStyle.set(qn("w:val"), "TableGrid")
    tblPr.insert(0, tblStyle)


def _suppress_tbl_look(tbl) -> None:
    """
    Zero out tblLook so Word doesn't apply built-in style overrides that
    swallow custom border definitions.
    """
    tblPr = _get_or_create_tblPr(tbl)
    for old in tblPr.findall(qn("w:tblLook")):
        tblPr.remove(old)
    look = OxmlElement("w:tblLook")
    look.set(qn("w:val"), "0000")
    look.set(qn("w:firstRow"), "0")
    look.set(qn("w:lastRow"), "0")
    look.set(qn("w:firstColumn"), "0")
    look.set(qn("w:lastColumn"), "0")
    look.set(qn("w:noHBand"), "1")
    look.set(qn("w:noVBand"), "1")
    tblPr.append(look)


def _set_tbl_grid_borders(tbl, color: str = COL_DIVIDER, size: str = "8") -> None:
    tblPr = _get_or_create_tblPr(tbl)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def _remove_tbl_borders(tbl) -> None:
    tblPr = _get_or_create_tblPr(tbl)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tbl_borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:color"), "auto")
        tbl_borders.append(el)
    tblPr.append(tbl_borders)


def _set_col_width(cell, width_dxa: int) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _cell_shading(cell, fill_hex: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _cell_borders_none(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:color"), "auto")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _cell_borders(cell, color: str = COL_DIVIDER, size: str = "8") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _cell_valign(cell, val: str = "center") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(old)
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), val)
    tcPr.append(vAlign)


def _cell_margins(cell, top: int = 80, bottom: int = 80, left: int = 120, right: int = 120) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _para_spacing(para, before: int = 0, after: int = 0) -> None:
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    pPr.append(spacing)


# ── Styles / Document setup ────────────────────────────────────────────────────

def _ensure_styles(doc: Document) -> None:
    styles = doc.styles

    def _get_or_create(name: str):
        try:
            return styles[name]
        except KeyError:
            return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.5)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    section = _get_or_create("PdfSectionTitle")
    section.font.name = "Calibri"
    section.font.size = Pt(10.5)
    section.font.bold = True
    section.font.color.rgb = RGBColor.from_string(COL_SECTION_FG)

    group = _get_or_create("PdfGroupTitle")
    group.font.name = "Calibri"
    group.font.size = Pt(9.5)
    group.font.bold = True
    group.font.color.rgb = RGBColor.from_string(COL_SECTION_FG)

    muted = _get_or_create("PdfMuted")
    muted.font.name = "Calibri"
    muted.font.size = Pt(8)
    muted.font.color.rgb = RGBColor.from_string(COL_MUTED)


def _new_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(0.8)
    _ensure_styles(doc)
    return doc


# ── Word page header (printed on every page) ───────────────────────────────────

def _build_page_header(doc: Document) -> None:
    """
    Logo + pharmacy name + email in the Word header section.
    Logo column is kept narrow (just wider than the logo image) so the text
    sits close beside it.
    """
    logo_path = _static_abs_path("img/app_icon_trans-512x512.png")
    header = doc.sections[0].header

    # Remove the default empty paragraph Word inserts into every header
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    # Newer python-docx requires width when calling add_table on a BlockItemContainer
    tbl = header.add_table(rows=1, cols=2, width=Twips(DOC_WIDTH_DXA))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_tbl_width(tbl, DOC_WIDTH_DXA)
    _set_tbl_fixed_layout(tbl)
    _set_tbl_grid(tbl, [HEADER_LOGO_COL_DXA, HEADER_INFO_COL_DXA])
    _remove_tbl_borders(tbl)
    _suppress_tbl_look(tbl)

    logo_cell = tbl.rows[0].cells[0]
    info_cell = tbl.rows[0].cells[1]

    _set_col_width(logo_cell, HEADER_LOGO_COL_DXA)
    _set_col_width(info_cell, HEADER_INFO_COL_DXA)

    for cell in (logo_cell, info_cell):
        _cell_borders_none(cell)

    _cell_margins(logo_cell, top=80, bottom=40, left=0, right=80)
    _cell_margins(info_cell, top=40, bottom=40, left=0, right=0)
    # Centre the text block vertically. Logo fills its cell so centering there has
    # no effect; centering only info_cell aligns both visual centres in the same row.
    _cell_valign(info_cell, "center")
    _cell_valign(logo_cell, "center")

    if logo_path and os.path.exists(logo_path):
        logo_para = logo_cell.paragraphs[0]
        _para_spacing(logo_para, 0, 0)
        logo_para.add_run().add_picture(logo_path, width=Cm(1.5))

    name_para = info_cell.paragraphs[0]
    _para_spacing(name_para, 0, 10)
    nr = name_para.add_run(PHARMACY_NAME)
    nr.bold = True
    nr.font.size = Pt(10)
    nr.font.color.rgb = RGBColor.from_string(COL_SECTION_FG)

    mail_para = info_cell.add_paragraph(style="PdfMuted")
    _para_spacing(mail_para, 0, 0)
    mail_para.add_run(PHARMACY_EMAIL)


# ── Document title block (body) ────────────────────────────────────────────────

def _add_logo_and_title(doc: Document, title: str, meta: dict, prepared_by, generated_at) -> None:
    """Document title, metadata and prepared-by — plain paragraphs."""
    p_title = doc.add_paragraph()
    _para_spacing(p_title, 0, 80)
    tr = p_title.add_run(title)
    tr.bold = True
    tr.font.size = Pt(16)
    tr.font.color.rgb = RGBColor.from_string(COL_SECTION_FG)

    for label, value in meta.items():
        p = doc.add_paragraph()
        _para_spacing(p, 0, 20)
        lr = p.add_run(f"{label}: ")
        lr.bold = True
        lr.font.size = Pt(9)
        vr = p.add_run(str(value))
        vr.font.size = Pt(9)

    name = f"{prepared_by.first_name} {prepared_by.last_name}".strip() or prepared_by.username

    p_prep = doc.add_paragraph()
    _para_spacing(p_prep, 0, 20)
    lr = p_prep.add_run("Voorbereid door: ")
    lr.bold = True
    lr.font.size = Pt(9)
    p_prep.add_run(name).font.size = Pt(9)

    p_date = doc.add_paragraph(style="PdfMuted")
    _para_spacing(p_date, 0, 160)
    p_date.add_run(f"Export: {generated_at.strftime('%d-%m-%Y  %H:%M')}")


# ── Section / structural helpers ──────────────────────────────────────────────

def _add_section_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="PdfSectionTitle")
    _para_spacing(p, 0, 60)
    p.add_run(text)


def _add_group_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="PdfGroupTitle")
    _para_spacing(p, 100, 16)
    p.paragraph_format.keep_with_next = True
    p.add_run(text)


def _add_patient_heading(doc: Document, naam: str, geboortedatum) -> None:
    p = doc.add_paragraph()
    _para_spacing(p, 160, 40)
    p.paragraph_format.keep_with_next = True

    name_run = p.add_run(naam)
    name_run.bold = True
    name_run.font.size = Pt(12)
    name_run.font.color.rgb = RGBColor.from_string(COL_SECTION_FG)

    dob = geboortedatum.strftime("%d-%m-%Y") if geboortedatum else "Onbekend"
    dob_run = p.add_run(f" ({dob})")
    dob_run.font.size = Pt(9)
    dob_run.font.color.rgb = RGBColor.from_string(COL_MUTED)


# ── Tables / Comment cards ─────────────────────────────────────────────────────

def _add_meds_table(doc: Document, meds: list) -> None:
    """
    Kolommen: Middel 35% | Gebruik 30% | Opmerking 35% van DOC_WIDTH_DXA.
    Header: lichtblauwe achtergrond, zwarte bold tekst.
    Datarijen: witte achtergrond met zichtbare grid-lijnen.
    """
    if not meds:
        return

    col_w = [3175, 2722, 3175]

    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_grid_style(tbl)
    _set_tbl_width(tbl, DOC_WIDTH_DXA)
    _set_tbl_grid_borders(tbl, color=COL_DIVIDER, size="8")  # before tblLayout (schema order)
    _set_tbl_fixed_layout(tbl)
    _set_tbl_grid(tbl, col_w)
    _suppress_tbl_look(tbl)

    for i, (cell, label) in enumerate(zip(tbl.rows[0].cells, ["Middel", "Gebruik", "Opmerking (Medimo)"])):
        _set_col_width(cell, col_w[i])
        _cell_borders(cell, color=COL_DIVIDER, size="8")  # before shd (schema order)
        _cell_shading(cell, COL_HEADER_BG)
        _cell_margins(cell, top=80, bottom=80, left=120, right=120)

        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    for gm in meds:
        row = tbl.add_row()

        clean = str(gm.get("clean", "") if isinstance(gm, dict) else getattr(gm, "clean", "") or "")
        gebruik = str(gm.get("gebruik", "") if isinstance(gm, dict) else getattr(gm, "gebruik", "") or "")
        opmerking = str(
            gm.get("opmerking", "") if isinstance(gm, dict) else getattr(gm, "opmerking", "") or ""
        ) or "-"

        for i, (cell, val) in enumerate(zip(row.cells, [clean, gebruik, opmerking])):
            _set_col_width(cell, col_w[i])
            _cell_borders(cell, color=COL_DIVIDER, size="8")  # before shd (schema order)
            _cell_shading(cell, "FFFFFF")
            _cell_margins(cell, top=80, bottom=80, left=120, right=120)

            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            if i == 0:
                run.bold = True


def _comment_card(doc: Document, label: str, lines: list) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_tbl_borders(tbl)
    _set_tbl_width(tbl, DOC_WIDTH_DXA)
    _set_tbl_fixed_layout(tbl)
    _set_tbl_grid(tbl, [DOC_WIDTH_DXA])
    _suppress_tbl_look(tbl)

    body = tbl.rows[0].cells[0]
    _set_col_width(body, DOC_WIDTH_DXA)
    _cell_borders_none(body)
    _cell_shading(body, COL_COMMENT_BG)
    _cell_margins(body, top=120, bottom=120, left=160, right=160)

    p0 = body.paragraphs[0]
    _para_spacing(p0, 0, 50)
    r0 = p0.add_run(label)
    r0.bold = True
    r0.font.size = Pt(8.5)
    r0.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    for i, line in enumerate(lines):
        p = body.add_paragraph()
        _para_spacing(p, 0, 0 if i < len(lines) - 1 else 30)
        p.add_run(line).font.size = Pt(8.5)


def _add_comment_box(doc: Document, label: str, text: str) -> None:
    if not text or not text.strip():
        return
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if not lines:
        return
    _comment_card(doc, label, lines)


def _add_empty_comment_box(doc: Document) -> None:
    _comment_card(doc, "Opmerkingen:", [""])


def _render_patient_block(doc: Document, block: PdfPatientBlock) -> None:
    for group_id, group_data in block.grouped_meds:
        _add_group_title(doc, group_data["naam"])

        meds = group_data.get("meds", [])
        _add_meds_table(doc, meds)

        comment = block.comments_lookup.get(group_id)
        historie_tekst = comment.historie if comment else ""
        opmerking_tekst = comment.tekst if comment else ""

        if historie_tekst:
            _add_comment_box(doc, "Eerder besproken", historie_tekst)

        if opmerking_tekst:
            _add_comment_box(doc, "Opmerkingen:", opmerking_tekst)
        else:
            _add_empty_comment_box(doc)


# ── Responses ──────────────────────────────────────────────────────────────────

def _build_response(doc: Document, filename: str) -> HttpResponse:
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    resp = HttpResponse(
        buffer.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    resp["Content-Disposition"] = f'attachment; filename="{filename}"'
    return resp


def _build_patient_docx_response(patient: MedicatieReviewPatient, user) -> HttpResponse:
    block = _build_patient_block(patient)
    doc = _new_doc()
    now = timezone.localtime(timezone.now())
    dob = patient.geboortedatum.strftime("%d-%m-%Y") if patient.geboortedatum else "Onbekend"

    _build_page_header(doc)
    _add_logo_and_title(
        doc,
        "Medicatiebeoordeling",
        {
            "Patiënt": patient.naam,
            "Geboortedatum": dob,
            "Afdeling": block.afdeling_naam or "-",
        },
        user,
        now,
    )

    _add_section_title(doc, "Medicatieoverzicht & Opmerkingen")
    _render_patient_block(doc, block)

    safe_name = (patient.naam or "patient").replace("/", "-")
    return _build_response(doc, f"medicatiebeoordeling_{safe_name}.docx")


# ── Views ──────────────────────────────────────────────────────────────────────

@login_required
def export_patient_review_docx(request, pk: int) -> HttpResponse:
    if not can(request.user, "can_view_medicatiebeoordeling"):
        return HttpResponseForbidden()

    patient = (
        MedicatieReviewPatient.objects.select_related("afdeling")
        .prefetch_related("comments", "med_group_overrides")
        .filter(pk=pk)
        .first()
    )
    if not patient:
        raise Http404()

    return _build_patient_docx_response(patient, request.user)


@login_required
def export_afdeling_review_docx(request, pk: int) -> HttpResponse:
    if not can(request.user, "can_view_medicatiebeoordeling"):
        return HttpResponseForbidden()

    afdeling = MedicatieReviewAfdeling.objects.filter(pk=pk).first()
    if not afdeling:
        raise Http404()

    patienten = (
        afdeling.patienten.all()
        .select_related("afdeling")
        .prefetch_related("comments", "med_group_overrides")
        .order_by("naam")
    )
    blocks = [_build_patient_block(p) for p in patienten]
    doc = _new_doc()
    now = timezone.localtime(timezone.now())

    _build_page_header(doc)
    _add_logo_and_title(
        doc,
        "Medicatiebeoordeling",
        {"Afdeling": afdeling.afdeling},
        request.user,
        now,
    )

    _add_section_title(doc, "Overzicht patiënten")

    col_w_patients = [5436, 3636]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_grid_style(tbl)
    _set_tbl_width(tbl, DOC_WIDTH_DXA)
    _set_tbl_grid_borders(tbl, color=COL_DIVIDER, size="8")  # before tblLayout (schema order)
    _set_tbl_fixed_layout(tbl)
    _set_tbl_grid(tbl, col_w_patients)
    _suppress_tbl_look(tbl)

    for i, (cell, label) in enumerate(zip(tbl.rows[0].cells, ["Naam", "Geboortedatum"])):
        _set_col_width(cell, col_w_patients[i])
        _cell_borders(cell, color=COL_DIVIDER, size="8")  # before shd (schema order)
        _cell_shading(cell, COL_HEADER_BG)
        _cell_margins(cell, top=80, bottom=80, left=120, right=120)

        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    for block in blocks:
        row = tbl.add_row()
        dob = block.patient.geboortedatum.strftime("%d-%m-%Y") if block.patient.geboortedatum else "Onbekend"

        for i, (cell, val) in enumerate(zip(row.cells, [block.patient.naam, dob])):
            _set_col_width(cell, col_w_patients[i])
            _cell_borders(cell, color=COL_DIVIDER, size="8")  # before shd (schema order)
            _cell_shading(cell, "FFFFFF")
            _cell_margins(cell, top=80, bottom=80, left=120, right=120)

            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            if i == 0:
                run.bold = True

    doc.add_page_break()
    for i, block in enumerate(blocks):
        if i > 0:
            doc.add_page_break()

        _add_patient_heading(doc, block.patient.naam, block.patient.geboortedatum)
        _render_patient_block(doc, block)

    safe_name = (afdeling.afdeling or "afdeling").replace("/", "-")
    return _build_response(doc, f"medicatiebeoordeling_{safe_name}.docx")